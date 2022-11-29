from prof_risks.DBAdapter import DBAdapter

from docx import Document

from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Mm, Cm, Pt

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.shared import OxmlElement, qn


class DocxAdapter:
    def __init__(self, organization_name, lab, contract, date_otch):
        self.db = DBAdapter()
        self.organization_name = organization_name
        self.lab = lab
        self.contract = contract
        self.date_otch = date_otch
        self.cards_doc = self.create_cards_file()

    def generate_all(self):
        self.generate_cards()
        self.generate_otchet_file()

    def generate_cards(self):
        # doc = self.create_cards_file()
        departments = self.db.get_departments_from_DB(self.organization_name)
        count_cards = 1
        for dep in departments:
            work_places = self.db.get_work_places_from_DB(dep[0])
            for wp in work_places:
                if count_cards != 1:
                    self.cards_doc.add_page_break()
                dangers = self.db.get_dangers_from_DB(wp[0])
                persons = self.db.get_persons_from_DB(wp[0])
                self.add_card_in_doc(dep[1], wp[1], wp[2], dangers, count_cards, persons)
                count_cards += 1

        file_name = f'output/Риски/{self.organization_name}/карты.docx'
        self.cards_doc.save(file_name)

    def add_card_in_doc(self, department, wp_caption, wp_equipment, dangers, card_number, persons):
        print(card_number)
        # Add card head
        table1 = self.cards_doc.add_table(rows=2, cols=3)
        table1_cells = table1.rows[0].cells
        table1.autofit = True
        table1.cell(0, 0).merge(table1.cell(1, 0)).paragraphs[0].add_run().add_picture(
            f'input/for_risks/logo/{self.lab}.jpg',
            width=Mm(30))
        run1 = table1.cell(0, 1).paragraphs[0].add_run("Карта оценки уровней профессионального риска")
        run1.font.size = Pt(10)
        run1.bold = True
        table1.cell(0, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table1.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run2 = table1.cell(1, 1).paragraphs[0].add_run(f"№   {card_number}   ")
        run2.font.size = Pt(10)
        # run2.font.underline = True
        run2.bold = True
        table1.cell(1, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table1.cell(1, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run3 = table1.cell(0, 2).merge(table1.cell(1, 2)).paragraphs[0].add_run(self.contract)
        run3.font.size = Pt(8)
        run3.bold = True
        table1.cell(0, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        table1.cell(0, 2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        table1.columns[0].width = Mm(30)
        table1_cells[0].width = Mm(30)
        table1.columns[1].width = Mm(175)
        table1_cells[1].width = Mm(175)
        table1.columns[2].width = Mm(55)
        table1_cells[2].width = Mm(55)

        self.cards_doc.add_paragraph("")

        # Add table2
        table2 = self.cards_doc.add_table(rows=3, cols=2)

        table2.autofit = True
        table2.style = 'Table Grid'

        table2.cell(0, 0).paragraphs[0].add_run("Наименование профессии / должности").font.size = Pt(8)
        table2.cell(1, 0).paragraphs[0].add_run("Наименование подразделения").font.size = Pt(8)
        table2.cell(2, 0).paragraphs[0].add_run("Применяемое оборудование").font.size = Pt(8)

        table2.cell(0, 1).paragraphs[0].add_run(wp_caption).font.size = Pt(8)
        table2.cell(1, 1).paragraphs[0].add_run(department).font.size = Pt(8)
        table2.cell(2, 1).paragraphs[0].add_run(wp_equipment).font.size = Pt(8)
        table2.columns[0].width = Mm(50)
        table2.columns[1].width = Mm(210)
        for i in range(0, 3):
            table2_cells = table2.rows[i].cells
            table2_cells[0].width = Mm(50)
            table2_cells[1].width = Mm(210)
            table2_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            table2_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table2_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table2_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        self.cards_doc.add_paragraph("")
        # Add table3
        table3 = self.cards_doc.add_table(rows=3, cols=8)

        table3.autofit = True
        table3.style = 'Table Grid'

        table3.cell(0, 0).merge(table3.cell(2, 0)).paragraphs[0] \
            .add_run("Идентификация опасностей\n(код и наименование опасности)").font.size = Pt(8)

        table3.cell(0, 1).merge(table3.cell(0, 5)).paragraphs[0].add_run("Индекс профессионального риска (ИПР)") \
            .font.size = Pt(8)

        table3.cell(1, 1).merge(table3.cell(1, 2)).paragraphs[0].add_run("Вероятность") \
            .font.size = Pt(8)
        table3.cell(1, 3).merge(table3.cell(1, 4)).paragraphs[0].add_run("Тяжесть ущерба") \
            .font.size = Pt(8)

        table3.cell(2, 1).paragraphs[0].add_run("Оценка").font.size = Pt(8)
        table3.cell(2, 2).paragraphs[0].add_run("Балл").font.size = Pt(8)
        table3.cell(2, 3).paragraphs[0].add_run("Оценка").font.size = Pt(8)
        table3.cell(2, 4).paragraphs[0].add_run("Балл").font.size = Pt(8)

        table3.cell(1, 5).merge(table3.cell(2, 5)).paragraphs[0].add_run("Итог").font.size = Pt(8)

        table3.cell(0, 6).merge(table3.cell(2, 6)).paragraphs[0].add_run("Комментарии аудитора").font.size = Pt(8)
        table3.cell(0, 7).merge(table3.cell(2, 7)).paragraphs[0].add_run("Корректирующие мероприятия").font.size = Pt(8)

        for row in table3.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        table3.columns[0].width = Mm(55)
        table3.rows[0].cells[0].width = Mm(55)
        table3.rows[0].cells[1].width = Mm(134)
        table3.columns[6].width = Mm(42)
        table3.rows[0].cells[6].width = Mm(42)
        table3.columns[7].width = Mm(34)
        table3.rows[0].cells[7].width = Mm(34)

        probability_list = ['', 'Почти невозможно', 'Маловероятно', 'Может быть', 'Вероятно',
                            'Почти наверняка']
        severity_list = ['', 'Незначительная', 'Низкая', 'Средняя', 'Значительная', 'Катастрофическая']

        for danger in dangers:
            cells3 = table3.add_row().cells
            grade = danger[1] * danger[2]
            res = [danger[0], probability_list[danger[1]], danger[1], severity_list[danger[2]],
                   danger[2], grade, danger[3], danger[4]]
            for i, item in enumerate(res):
                cells3[i].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells3[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells3[i].paragraphs[0].add_run(str(item)).font.size = Pt(8)

            colors = [r'<w:shd {} w:fill="#92D050"/>', r'<w:shd {} w:fill="#FFFF00"/>',
                      r'<w:shd {} w:fill="#FF0000"/>']
            result_str = ''
            if grade < 5:
                result_str = colors[0]
            else:
                if grade < 12:
                    result_str = colors[1]
                else:
                    result_str = colors[2]

            shading_elm = parse_xml(result_str.format(nsdecls('w')))
            cells3[5]._tc.get_or_add_tcPr().append(shading_elm)

        self.cards_doc.add_paragraph("")
        self.fill_signature(0, 'Инженер по специальной оценке условий труда ИЛ', ['Василенко А.С.', ], self.date_otch)
        self.fill_signature(1, wp_caption, persons, '')
        return

    def fill_signature(self, status, position, names, date):
        status_list = ["Разработал:", "Ознакомлен:"]
        run = self.cards_doc.add_paragraph().add_run(status_list[status])
        run.bold = True
        run.font.size = Pt(8)
        table = self.cards_doc.add_table(rows=0, cols=7)
        names_list = []
        if status != 1:
            for name in names:
                names_list.append(name)
        if status == 1:
            for i in range(10):
                names_list.append('')
        for name in names_list:
            row_1 = table.add_row()
            row_2 = table.add_row()

            for i in [0, 2, 4, 6]:
                cell = row_1.cells[i]
                bottom_border = OxmlElement('w:bottom')
                bottom_border.set(qn('w:val'), 'single')
                bottom_border.set(qn('w:sz'), '4')
                cell._tc.get_or_add_tcPr().append(bottom_border)

            row_1.cells[0].paragraphs[0].add_run(position if name != '' else '').font.size = Pt(8)
            row_2.cells[0].paragraphs[0].add_run("(должность)").font.size = Pt(6)
            row_1.cells[2].paragraphs[0].add_run(name).font.size = Pt(8)
            row_2.cells[2].paragraphs[0].add_run("(Ф.И.О.)").font.size = Pt(6)
            row_1.cells[4].paragraphs[0].add_run(date if name != '' else '').font.size = Pt(8)
            row_2.cells[4].paragraphs[0].add_run("(Дата)").font.size = Pt(6)
            row_1.cells[6].paragraphs[0].add_run('').font.size = Pt(8)
            row_2.cells[6].paragraphs[0].add_run("(Подпись)").font.size = Pt(6)

            table.columns[0].width = Mm(65)
            row_1.cells[0].width = Mm(65)
            table.columns[1].width = Mm(20)
            row_1.cells[1].width = Mm(20)
            table.columns[2].width = Mm(65)
            row_1.cells[2].width = Mm(65)
            table.columns[3].width = Mm(20)
            row_1.cells[3].width = Mm(20)
            table.columns[4].width = Mm(40)
            row_1.cells[4].width = Mm(40)
            table.columns[5].width = Mm(20)
            row_1.cells[5].width = Mm(20)
            table.columns[6].width = Mm(40)
            row_1.cells[6].width = Mm(40)
            for i in range(7):
                cell = row_1.cells[i]
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

                cell2 = row_2.cells[i]
                cell2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    def generate_otchet_file(self):
        doc = Document(f'C:\\Users\\buhu_\\PycharmProjects\\OPR\\input\\for_risks\\templates\\{self.lab}.docx')
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        all_tables = doc.tables
        self.fill_table_1(all_tables[10])
        self.fill_table_2(all_tables[11])
        file_name = f'output/Риски/{self.organization_name}/отчет.docx'
        doc.save(file_name)

    def fill_table_1(self, table):
        count = 1
        table.style = 'Table Grid'
        res_table_1 = self.db.get_table_1_data(self.organization_name)
        probability_list = ['', 'Почти\nневоз-\nможно', 'Малове-\nроятно', 'Может\nбыть', 'Веро-\nятно',
                            'Почти\nнавер-\nняка']
        severity_list = ['', 'Незна-\nчитель-\nная', 'Низкая', 'Сред-\nняя', 'Значи-\nтельная',
                         'Ката-\nстрофи-\nческая']
        colors = [r'<w:shd {} w:fill="#92D050"/>', r'<w:shd {} w:fill="#FFFF00"/>',
                  r'<w:shd {} w:fill="#FF0000"/>']
        for row in res_table_1:
            print(f'{count}')
            cells = table.add_row().cells
            res_row = [count, row[0] + ' / ' + row[1], row[7], row[2], probability_list[row[3]], row[3],
                       severity_list[row[4]], row[4], row[3] * row[4], row[5], row[6]]
            for i, item in enumerate(res_row):
                cells[i].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[i].paragraphs[0].add_run(str(item)).font.size = Pt(9)

            shading_elm = parse_xml(colors[self.get_color_index(cells[8].paragraphs[0].text)].format(nsdecls('w')))
            cells[8]._tc.get_or_add_tcPr().append(shading_elm)
            count += 1

    def fill_table_2(self, table2):
        count = 1
        table2.style = 'Table Grid'
        res_table_2 = self.db.get_table_2_data(self.organization_name)
        colors = [r'<w:shd {} w:fill="#92D050"/>', r'<w:shd {} w:fill="#FFFF00"/>',
                  r'<w:shd {} w:fill="#FF0000"/>']
        cell4_text = ['Низкий', 'Умеренный', 'Высокий']
        for row in res_table_2:
            cells2 = table2.add_row().cells
            cell3_text = f'Опасность:\n\n{row[2]}\n\nКоментарии аудитора:\n\n{row[5]}'
            res = [count, row[0] + ' / ' + row[1], cell3_text, cell4_text[self.get_color_index(row[3] * row[4])],
                   row[6], row[3], row[3], row[4], row[4], row[3] * row[4], row[3] * row[4]]
            for i, item in enumerate(res):
                cells2[i].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells2[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells2[i].paragraphs[0].add_run(str(item)).font.size = Pt(9)

            shading_elm = parse_xml(colors[self.get_color_index(cells2[9].paragraphs[0].text)].format(nsdecls('w')))
            cells2[9]._tc.get_or_add_tcPr().append(shading_elm)

            shading_elm2 = parse_xml(colors[self.get_color_index(cells2[10].paragraphs[0].text)].format(nsdecls('w')))
            cells2[10]._tc.get_or_add_tcPr().append(shading_elm2)
            count += 1

        return 0

    @staticmethod
    def get_color_index(number):
        if int(number) < 5:
            return 0
        elif int(number) < 12:
            return 1
        else:
            return 2

    @staticmethod
    def create_cards_file():
        # Create and customization document
        doc = Document()
        current_section = doc.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        current_section.orientation = WD_ORIENTATION.LANDSCAPE
        current_section.page_width = new_width
        current_section.page_height = new_height
        current_section.left_margin = Cm(2)
        current_section.right_margin = Cm(2)
        current_section.top_margin = Cm(2)
        current_section.bottom_margin = Mm(15)
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        return doc
