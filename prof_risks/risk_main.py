import csv
from sqlalchemy import create_engine, select, MetaData, Table, insert, text

from docx.enum.section import WD_ORIENTATION
from docx import Document
from docx.shared import Mm, Cm, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import ns
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement, qn
import json


def import_factors_from_csv(csv_file):
    engine = create_engine('sqlite:///prof_risks/sqlite.db')
    metadata = MetaData(engine)
    conn = engine.connect()
    dangers_type = Table('dangers_type', metadata, autoload=True)
    dangers_dict = Table('dangers_dict', metadata, autoload=True)

    with open(csv_file, newline='', encoding="utf-8-sig") as File:
        reader = csv.reader(File, delimiter=';')
        for row in reader:
            risk_type = conn.execute(select([dangers_type.c.id]).where(dangers_type.c.name == row[1])).fetchone()
            if risk_type is None:
                risk_type = conn.execute(insert(dangers_type).values(name=row[1])).inserted_primary_key
            conn.execute(insert(dangers_dict).values(id=row[0], description=row[2], type=risk_type[0]))


def import_organization_csv(csv_file, org_name, org_head_pos='', org_head_name=''):
    engine = create_engine('sqlite:///prof_risks/sqlite.db')
    metadata = MetaData(engine)
    conn = engine.connect()
    organizations = Table('organizations', metadata, autoload=True)
    departments = Table('departments', metadata, autoload=True)
    work_places = Table('work_places', metadata, autoload=True)
    result = Table('result', metadata, autoload=True)
    dangers_group = Table('dangers_group', metadata, autoload=True)
    insert_org_sql = insert(organizations).values(name=org_name, head_position=org_head_pos, head_name=org_head_name)
    org_id = conn.execute(insert_org_sql).inserted_primary_key[0]
    with open(csv_file, newline='', encoding="utf-8-sig") as File:
        reader = csv.reader(File, delimiter=';')
        for row in reader:
            dep_id = conn.execute(select([departments.c.id]).where(departments.c.name == row[1])).fetchone()
            if dep_id is None:
                insert_dep_sql = insert(departments).values(name=row[1], address='', organization_id=org_id)
                dep_id = conn.execute(insert_dep_sql).inserted_primary_key
            insert_wp = insert(work_places).values(name=row[0], equipment=row[2], department_id=dep_id[0])
            wp_id = conn.execute(insert_wp).inserted_primary_key[0]
            wp_template = conn.execute(select([dangers_group]).where(dangers_group.c.name == row[7]))
            for risk in wp_template:
                insert_result_sql = insert(result).values(work_place_id=wp_id, danger_id=risk[2], probability=risk[3],
                                                          severity=risk[4], comment=risk[5], measures=risk[6])
                conn.execute(insert_result_sql)


def import_organization_json(json_file, org_name, org_head_pos='', org_head_name=''):
    engine = create_engine('sqlite:///prof_risks/sqlite.db')
    metadata = MetaData(engine)
    conn = engine.connect()
    organizations = Table('organizations', metadata, autoload=True)
    departments = Table('departments', metadata, autoload=True)
    work_places = Table('work_places', metadata, autoload=True)
    result = Table('result', metadata, autoload=True)
    dangers_group = Table('dangers_group', metadata, autoload=True)
    persons = Table('persons', metadata, autoload=True)
    insert_org_sql = insert(organizations).values(name=org_name, head_position=org_head_pos, head_name=org_head_name)
    org_id = conn.execute(insert_org_sql).inserted_primary_key[0]
    rm_dict = {}
    with open(json_file, 'r', encoding="utf-8-sig") as file:
        rm_dict = json.load(file)
        for ceh in rm_dict['ceh'].keys():
            add_uch(ceh, rm_dict['ceh'][ceh]['uch'], org_id, conn, departments, work_places, dangers_group, result,
                    persons)


def add_uch(curr_uch, uchs, org_id, conn, departments, work_places, dangers_group, result, persons):
    for uch in uchs.keys():
        if uch == 'rm':
            dep_id = insert_department(curr_uch, org_id, conn, departments)
            for rm in uchs['rm']:
                insert_work_place(rm, dep_id[0], conn, work_places, dangers_group, result, persons)
        else:
            curr_uch += ' - ' + uch
            add_uch(curr_uch, uchs[uch], org_id, conn, departments, work_places, dangers_group, result, persons)


def insert_department(department_name, org_id, conn, departments):
    dep_id = conn.execute(select([departments.c.id]).where(departments.c.name == department_name)).fetchone()
    if dep_id is None:
        insert_dep_sql = insert(departments).values(name=department_name, address='', organization_id=org_id)
        dep_id = conn.execute(insert_dep_sql).inserted_primary_key
    return dep_id


def insert_work_place(wp, dep_id, conn, work_places, dangers_group, result, persons):
    insert_wp = insert(work_places).values(name=wp['caption'], equipment=wp['oborud'], department_id=dep_id)
    wp_id = conn.execute(insert_wp).inserted_primary_key[0]
    for name in wp['fio']:
        insert_fio = insert(persons).values(work_place_id=wp_id, name=name)
        conn.execute(insert_fio)
    wp_template = conn.execute(select([dangers_group]).where(dangers_group.c.name == wp['rm_type']))
    for risk in wp_template:
        insert_result_sql = insert(result).values(work_place_id=wp_id, danger_id=risk[2], probability=risk[3],
                                                  severity=risk[4], comment=risk[5], measures=risk[6])
        conn.execute(insert_result_sql)


def create_cards_docx(organization_name):
    # Create and customization document
    doc = Document()
    current_section = doc.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    current_section.orientation = WD_ORIENTATION.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height
    current_section.left_margin = Cm(2)
    current_section.right_margin = Cm(2)
    current_section.top_margin = Cm(3)
    current_section.bottom_margin = Mm(15)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    doc.add_paragraph().add_run("Раздел 4. Оценка профессиональных рисков").bold = True
    table = doc.add_table(rows=3, cols=11)
    table.autofit = True
    table.style = 'Table Grid'
    table_cells = table.rows[0].cells
    table.columns[0].width = Mm(8.9)
    table.columns[1].width = Mm(33.5)
    table.columns[2].width = Mm(27.5)
    table.columns[3].width = Mm(47.5)
    table.columns[9].width = Mm(50.2)
    table.columns[10].width = Mm(32.4)
    table_cells[0].width = Mm(10)
    table.cell(0, 0).merge(table.cell(2, 0)).paragraphs[0].add_run("№\nп/п").font.size = Pt(9)
    table.cell(0, 1).merge(table.cell(2, 1)).paragraphs[0].add_run("Наименование\nпрофессии").font.size = Pt(9)
    table.cell(0, 2).merge(table.cell(2, 2)).paragraphs[0].add_run("Объект оценки\nпрофессиональных\nрисков") \
        .font.size = Pt(9)
    table.cell(0, 3).merge(table.cell(2, 3)).paragraphs[0].add_run("Идентификация\nопасностей (код "
                                                                   "и\nнаименование\nопасности)").font.size = Pt(9)
    table.cell(0, 4).merge(table.cell(0, 8)).paragraphs[0].add_run("Индекс профессионального риска (ИПР)") \
        .font.size = Pt(9)
    table.cell(1, 4).merge(table.cell(1, 5)).paragraphs[0].add_run("Вероятность").font.size = Pt(9)
    table.cell(1, 6).merge(table.cell(1, 7)).paragraphs[0].add_run("Тяжесть ущерба").font.size = Pt(9)
    table.cell(2, 4).paragraphs[0].add_run("Оценка").font.size = Pt(9)
    table.cell(2, 5).paragraphs[0].add_run("Балл").font.size = Pt(9)
    table.cell(2, 6).paragraphs[0].add_run("Оценка").font.size = Pt(9)
    table.cell(2, 7).paragraphs[0].add_run("Балл").font.size = Pt(9)
    table.cell(1, 8).merge(table.cell(2, 8)).paragraphs[0].add_run("Итог").font.size = Pt(9)
    table.cell(0, 9).merge(table.cell(2, 9)).paragraphs[0].add_run("Комментарии\nаудитора").font.size = Pt(9)
    table.cell(0, 10).merge(table.cell(2, 10)).paragraphs[0].add_run("Корректирующие\nмероприятия").font.size = Pt(9)

    for i in range(0, 3):
        for j in range(0, 11):
            cell = table.cell(i, j)
            cell.paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shading_elm = parse_xml(r'<w:shd {} w:fill="DEEAF6"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
    fill_table_risk_evaluation(table, organization_name)
    file_name = f'Output/test_cards.docx'
    doc.save(file_name)


def create_docx_from_template(organization_name):
    doc = Document('C:\\Users\\buhu_\\PycharmProjects\\OPR\\input\\for_risks\\templates\\prommash.docx')
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    all_tables = doc.tables
    fill_table_risk_evaluation(all_tables[10], all_tables[11], organization_name)
    file_name = f'Output/test_cards2.docx'
    doc.save(file_name)


def fill_table_risk_evaluation(table, table2, org_name):
    engine = create_engine('sqlite:///prof_risks/sqlite.db')
    conn = engine.connect()
    sql = text("""SELECT d.name, wp.name, dangers_dict.description, r.probability, r.severity, r.comment, r.measures 
    FROM (result r JOIN dangers_dict ON r.danger_id = dangers_dict.id) JOIN work_places wp ON wp.id = r.work_place_id 
    JOIN departments d ON wp.department_id = d.id JOIN organizations o ON d.organization_id = o.id 
    WHERE o.name = :name""")
    dic = {'name': org_name}
    res = conn.execute(sql, **dic)
    probability_list = ['', 'Почти\nневоз-\nможно', 'Малове-\nроятно', 'Может\nбыть', 'Веро-\nятно',
                        'Почти\nнавер-\nняка']
    severity_list = ['', 'Ката-\nстрофи-\nческая', 'Значи-\nтельная', 'Сред-\nняя', 'Низкая', 'Незна-\nчитель-\nная']
    count = 1
    table.style = 'Table Grid'
    table2.style = 'Table Grid'
    for row in res:
        cells = table.add_row().cells
        res = [count, row[0] + ' / ' + row[1], 'Процесс', row[2], probability_list[row[3]], row[3], severity_list[row[4]],
               row[4], row[3] * row[4], row[5], row[6]]
        for i, item in enumerate(res):
            cells[i].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(str(item)).font.size = Pt(9)
        count += 1
        colors = [r'<w:shd {} w:fill="#92D050"/>', r'<w:shd {} w:fill="#FFFF00"/>',
                  r'<w:shd {} w:fill="#FF0000"/>']
        result_str = ''
        cell4_text = 'Низкий'
        if row[3] * row[4] < 5:
            result_str = colors[0]
        else:
            if row[3] * row[4] < 12:
                result_str = colors[1]
                cell4_text = 'Умеренный'
            else:
                result_str = colors[2]
                cell4_text = 'Высокий'

            cells2 = table2.add_row().cells
            cell3_text = f'Опасность:\n\n{row[2]}\n\nКоментарии аудитора:\n\n{row[5]}'
            res = [count, row[0] + ' / ' + row[1], cell3_text, cell4_text, row[6], row[3], row[3], row[4], row[4],
                   row[3] * row[4], row[3] * row[4]]
            for i, item in enumerate(res):
                cells2[i].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells2[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells2[i].paragraphs[0].add_run(str(item)).font.size = Pt(9)
            shading_elm = parse_xml(result_str.format(nsdecls('w')))
            cells2[9]._tc.get_or_add_tcPr().append(shading_elm)
            shading_elm2 = parse_xml(result_str.format(nsdecls('w')))
            cells2[10]._tc.get_or_add_tcPr().append(shading_elm2)
        shading_elm = parse_xml(result_str.format(nsdecls('w')))
        cells[8]._tc.get_or_add_tcPr().append(shading_elm)
    return


def get_departments_from_DB(conn, org_name):
    sql = text("""SELECT d.id, d.name 
        FROM departments d JOIN organizations o ON d.organization_id = o.id 
        WHERE o.name = :name""")
    dic = {'name': org_name}
    departments = conn.execute(sql, **dic)
    return departments


def get_work_places_from_DB(conn, department_id):
    sql = text("""SELECT wp.id, wp.name, wp.equipment FROM work_places wp WHERE wp.department_id = :dep_id""")
    dic = {'dep_id': department_id}
    work_places = conn.execute(sql, **dic)
    return work_places


def get_dangers_from_DB(conn, wp_id):
    sql = text("""SELECT d.description, r.probability, r.severity, r.comment, r.measures 
        FROM result r JOIN dangers_dict d ON r.danger_id = d.id
        WHERE r.work_place_id = :wp_id""")
    dic = {'wp_id': wp_id}
    dangers = conn.execute(sql, **dic)
    return dangers


def get_persons_from_DB(conn, wp_id):
    sql = text("""SELECT p.name FROM persons p WHERE p.work_place_id = :wp_id""")
    dic = {'wp_id': wp_id}
    persons = conn.execute(sql, **dic)
    return persons


def generate_cards(organization_name, lab, contract):
    doc = create_cards_file()
    engine = create_engine('sqlite:///prof_risks/sqlite.db')
    conn = engine.connect()
    departments = get_departments_from_DB(conn, organization_name)
    count_cards = 1
    for dep in departments:
        work_places = get_work_places_from_DB(conn, dep[0])
        for wp in work_places:
            if count_cards != 1:
                doc.add_page_break()
            dangers = get_dangers_from_DB(conn, wp[0])
            persons = get_persons_from_DB(conn, wp[0])
            add_card_in_doc(doc, lab, contract, dep[1], wp[1], wp[2], dangers, count_cards, persons)
            count_cards += 1

    file_name = f'Output/test_cards.docx'
    doc.save(file_name)


def create_cards_file():
    # Create and customization document
    doc = Document()
    current_section = doc.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    current_section.orientation = WD_ORIENTATION.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height
    current_section.left_margin = Cm(2)
    current_section.right_margin = Cm(2)
    current_section.top_margin = Cm(2)
    current_section.bottom_margin = Mm(15)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    return doc


def fill_signature(doc, status, position, names, date):
    status_list = ["Разработал:", "Ознакомлен:"]
    run = doc.add_paragraph().add_run(status_list[status])
    run.bold = True
    run.font.size = Pt(8)
    table = doc.add_table(rows=0, cols=7)
    names_list = []
    for name in names:
        names_list.append(name)
    if status == 1:
        for i in range(10):
            names_list.append('')
    for name in names_list:
        row_1 = table.add_row()
        row_2 = table.add_row()

        for i in [0, 2, 4, 6]:
            cell = row_1.cells[i]
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '4')
            cell._tc.get_or_add_tcPr().append(bottom_border)

        row_1.cells[0].paragraphs[0].add_run(position if name != '' else '').font.size = Pt(8)
        row_2.cells[0].paragraphs[0].add_run("(должность)").font.size = Pt(6)
        row_1.cells[2].paragraphs[0].add_run(name).font.size = Pt(8)
        row_2.cells[2].paragraphs[0].add_run("(Ф.И.О.)").font.size = Pt(6)
        row_1.cells[4].paragraphs[0].add_run(date if name != '' else '').font.size = Pt(8)
        row_2.cells[4].paragraphs[0].add_run("(Дата)").font.size = Pt(6)
        row_1.cells[6].paragraphs[0].add_run('').font.size = Pt(8)
        row_2.cells[6].paragraphs[0].add_run("(Подпись)").font.size = Pt(6)

        table.columns[0].width = Mm(65)
        row_1.cells[0].width = Mm(65)
        table.columns[1].width = Mm(20)
        row_1.cells[1].width = Mm(20)
        table.columns[2].width = Mm(65)
        row_1.cells[2].width = Mm(65)
        table.columns[3].width = Mm(20)
        row_1.cells[3].width = Mm(20)
        table.columns[4].width = Mm(40)
        row_1.cells[4].width = Mm(40)
        table.columns[5].width = Mm(20)
        row_1.cells[5].width = Mm(20)
        table.columns[6].width = Mm(40)
        row_1.cells[6].width = Mm(40)
        for i in range(7):
            cell = row_1.cells[i]
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

            cell2 = row_2.cells[i]
            cell2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_card_in_doc(doc, lab, contract, department, wp_caption, wp_equipment, dangers, card_number, persons):
    # Add card head
    table1 = doc.add_table(rows=2, cols=3)
    table1_cells = table1.rows[0].cells
    table1.autofit = True
    table1.cell(0, 0).merge(table1.cell(1, 0)).paragraphs[0].add_run().add_picture(f'input/for_risks/logo/{lab}.jpg',
                                                                                   width=Mm(30))
    run1 = table1.cell(0, 1).paragraphs[0].add_run("Карта оценки уровней профессионального риска")
    run1.font.size = Pt(10)
    run1.bold = True
    table1.cell(0, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table1.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    run2 = table1.cell(1, 1).paragraphs[0].add_run(f"№   {card_number}   ")
    run2.font.size = Pt(10)
    # run2.font.underline = True
    run2.bold = True
    table1.cell(1, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table1.cell(1, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    run3 = table1.cell(0, 2).merge(table1.cell(1, 2)).paragraphs[0].add_run(contract)
    run3.font.size = Pt(8)
    run3.bold = True
    table1.cell(0, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    table1.cell(0, 2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    table1.columns[0].width = Mm(30)
    table1_cells[0].width = Mm(30)
    table1.columns[1].width = Mm(175)
    table1_cells[1].width = Mm(175)
    table1.columns[2].width = Mm(55)
    table1_cells[2].width = Mm(55)

    doc.add_paragraph("")

    # Add table2
    table2 = doc.add_table(rows=3, cols=2)

    table2.autofit = True
    table2.style = 'Table Grid'

    table2.cell(0, 0).paragraphs[0].add_run("Наименование профессии / должности").font.size = Pt(8)
    table2.cell(1, 0).paragraphs[0].add_run("Наименование подразделения").font.size = Pt(8)
    table2.cell(2, 0).paragraphs[0].add_run("Применяемое оборудование").font.size = Pt(8)

    table2.cell(0, 1).paragraphs[0].add_run(wp_caption).font.size = Pt(8)
    table2.cell(1, 1).paragraphs[0].add_run(department).font.size = Pt(8)
    table2.cell(2, 1).paragraphs[0].add_run(wp_equipment).font.size = Pt(8)
    table2.columns[0].width = Mm(50)
    table2.columns[1].width = Mm(210)
    for i in range(0, 3):
        table2_cells = table2.rows[i].cells
        table2_cells[0].width = Mm(50)
        table2_cells[1].width = Mm(210)
        table2_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        table2_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table2_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table2_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph("")
    # Add table3
    table3 = doc.add_table(rows=3, cols=8)

    table3.autofit = True
    table3.style = 'Table Grid'

    table3.cell(0, 0).merge(table3.cell(2, 0)).paragraphs[0]\
        .add_run("Идентификация опасностей\n(код и наименование опасности)").font.size = Pt(8)

    table3.cell(0, 1).merge(table3.cell(0, 5)).paragraphs[0].add_run("Индекс профессионального риска (ИПР)")\
        .font.size = Pt(8)

    table3.cell(1, 1).merge(table3.cell(1, 2)).paragraphs[0].add_run("Вероятность") \
        .font.size = Pt(8)
    table3.cell(1, 3).merge(table3.cell(1, 4)).paragraphs[0].add_run("Тяжесть ущерба") \
        .font.size = Pt(8)

    table3.cell(2, 1).paragraphs[0].add_run("Оценка").font.size = Pt(8)
    table3.cell(2, 2).paragraphs[0].add_run("Балл").font.size = Pt(8)
    table3.cell(2, 3).paragraphs[0].add_run("Оценка").font.size = Pt(8)
    table3.cell(2, 4).paragraphs[0].add_run("Балл").font.size = Pt(8)

    table3.cell(1, 5).merge(table3.cell(2, 5)).paragraphs[0].add_run("Итог").font.size = Pt(8)

    table3.cell(0, 6).merge(table3.cell(2, 6)).paragraphs[0].add_run("Комментарии аудитора").font.size = Pt(8)
    table3.cell(0, 7).merge(table3.cell(2, 7)).paragraphs[0].add_run("Корректирующие мероприятия").font.size = Pt(8)

    for row in table3.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    table3.columns[0].width = Mm(55)
    table3.rows[0].cells[0].width = Mm(55)
    table3.rows[0].cells[1].width = Mm(134)
    table3.columns[6].width = Mm(42)
    table3.rows[0].cells[6].width = Mm(42)
    table3.columns[7].width = Mm(34)
    table3.rows[0].cells[7].width = Mm(34)

    probability_list = ['', 'Почти невозможно', 'Маловероятно', 'Может быть', 'Вероятно',
                        'Почти наверняка']
    severity_list = ['', 'Катастрофическая', 'Значительная', 'Средняя', 'Низкая', 'Незначительная']

    for danger in dangers:
        cells3 = table3.add_row().cells
        grade = danger[1] * danger[2]
        res = [danger[0], probability_list[danger[1]], danger[1], severity_list[danger[2]],
               danger[2], grade, danger[3], danger[4]]
        for i, item in enumerate(res):
            cells3[i].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells3[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells3[i].paragraphs[0].add_run(str(item)).font.size = Pt(8)

        colors = [r'<w:shd {} w:fill="#92D050"/>', r'<w:shd {} w:fill="#FFFF00"/>',
                  r'<w:shd {} w:fill="#FF0000"/>']
        result_str = ''
        if grade < 5:
            result_str = colors[0]
        else:
            if grade < 12:
                result_str = colors[1]
            else:
                result_str = colors[2]

        shading_elm = parse_xml(result_str.format(nsdecls('w')))
        cells3[5]._tc.get_or_add_tcPr().append(shading_elm)

    doc.add_paragraph("")
    fill_signature(doc, 0, 'Инженер по специальной оценке условий труда ИЛ', ['Василенко А.С.', ], '20.09.2022')
    fill_signature(doc, 1, wp_caption, persons, '')

    return


def create_report_docx():
    pass
